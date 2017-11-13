#!/usr/bin/env python2
# -*- coding: utf-8 -*-
"""
Created on Tue Nov  7 05:48:32 2017

@author: ubuntu
"""

from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.shared import RGBColor
from docx.oxml.ns import qn

import os

document = Document()
#add heading
document.add_heading(u'Test',level=0)
document.add_heading(u'1.what?',1)
document.add_heading(u'2.why?',2)
document.add_heading(u'3.when?',3)
document.add_heading(u'4.where?',4)
document.add_heading(u'5.how?',5)
document.add_heading(u'6.whom?',6)
document.add_heading(u'7.whose?',7)
document.add_heading(u'8.which?',8)
document.add_heading(u'9.here?',9)
#add paragraph
paragraph = document.add_paragraph(u'I am on working')
#setting fontsize
run = paragraph.add_run(u'设置字号,')
run.font.size = Pt(24)
#set font
run = paragraph.add_run(u'Set Font,')
run.font.name = 'Consolas'
#set chinese font
run = paragraph.add_run(u'设置中文字体')
run.font.name = u'宋体'
run.font.size = 12
run.font.underline = True
run.font.bold = True
run.font.italic = True
#run.font.underline = WD_UNDERLINE.DOT_DASH
run.font.color.rgb =  RGBColor(0x42,0x24,0xE9)
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'),u'宋体')
#set italic
run = paragraph.add_run(u'斜体,')
run.italic = True
#set bold
run = paragraph.add_run(u'粗体')
run.bold = True
#insert quotation
document.add_paragraph('Intense quote',style='Intense Quote')
#insert
document.add_paragraph(u'non-ordered',style='List Bullet')
document.add_paragraph(u'ordered',style='List Number')
#add picture
document.add_picture(os.path.join(os.getcwd(),'book.jpg'),width=Inches(1.25))
#insert table
table = document.add_table(rows=1,cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Name'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
# insert tabel value
for i in xrange(3):
    row_cells = table.add_row().cells
    row_cells[0].text = 'test'+str(i)
    row_cells[1].text =  str(i)
    row_cells[2].text = 'desc'+str(i)
#insert sharing
document.add_page_break()
#save document
document.save(u'test.docx')