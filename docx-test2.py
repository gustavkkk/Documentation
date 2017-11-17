# -*- coding: utf-8 -*-
"""
Created on Mon Nov 13 09:30:33 2017

@author: Frank

ref:https://python-docx.readthedocs.io/en/latest/#api-documentation
    https://github.com/mikemaccana/python-docx/blob/master/example-makedocument.py
    https://stackoverflow.com/questions/25228106/how-to-extract-text-from-an-existing-docx-file-using-python-docx
    https://stackoverflow.com/questions/22765313/when-import-docx-in-python3-3-i-have-error-importerror-no-module-named-excepti
    https://stackoverflow.com/questions/24805671/how-to-use-python-docx-to-replace-text-in-a-word-document-and-save
    https://www.cnblogs.com/rencm/p/6285304.html
    
    docx/text/paragraph.py#Paragraph
    docx/blkcntnr.py#BlockItemContainer
    docx/oxml/text/paragraph.py#CT_P
    docx/oxml/document.py#CT_Body
    docx/oxml/xmlchemy.py#_OxmlElementBase--->all final operations based on lxml.etree.ElementBase functions
"""

import docx
import re
import jieba  
from config import dic

for word in dic:
    jieba.add_word(word)

def isChinese(text):
    ischinese = False
    if len(text)>1:
        for i in range(len(text)):
            if ord(text[i])>300:
                ischinese=True
                break
    else:
        if ord(text)>300:
            ischinese=True
    return ischinese

def remove_space(text,mode='all'):
    if mode == 'all':
        return ''.join(text.split())#text.replace(" ", "")
    elif mode == 'ending or leading':
        return text.strip()

class DOC:
    
    def __init__(self,filename=None):
        self.initialize()
        if filename is not None:
            self.load(filename)
        
    def load(self,filename='zhaobiao.docx'):
        self.doc = docx.Document(filename)
        
    def save(self,filename='output.docx'):
        self.doc.save(filename)
        
    def initialize(self):
        self.doc = None

    def getinfo(self):
        #print(len(self.doc.Pages))
        print(len(self.doc.paragraphs))
        print(len(self.doc.tables))
 
    def empty(self):
        self.doc._body.clear_content()
       
    def processtext(self):
        fullText = []
        paragraphs_ = []
        for i,paragraph in enumerate(self.doc.paragraphs):
            if len(paragraph.text) != 0:
                #print(i,paragraph.text)
                #style = paragraph.style
                #text = paragraph.text
                ptext = remove_space(paragraph.text)
                seg_list = jieba.cut(ptext, cut_all=False)
                DOC.deleteparagraph(paragraph)
                print("Default Mode: " + "/".join(seg_list))
                '''
                inline = paragraph.runs
                for j in range(len(inline)):
                    text = inline[j].text
                    if isChinese(text):
                        print(i,j,text)
                        #inline[j].text = "A"
                '''
            #fullText.append(paragraph.text)
        #self.doc._body._element.p_lst = []
        #self.doc._body._element.tbl_lst
        return '\n'.join(fullText)

    @staticmethod
    def deletepage(paragraph):
        pass
    
    @staticmethod
    def deleteparagraph(paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None

    @property
    def tables(self):
        return self.doc.tables
    
    @property
    def paragraphs(self):
        return self.doc.paragraphs
        
    def processtable(self):
        for i,table in enumerate(self.doc.tables):
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        print(paragraph.text)
                        if 'sea' in paragraph.text:
                            paragraph.text = 'ocean'
    
    def replace_regex(self, regex=re.compile(r"sea") , replace=r"ocean"):
    
        for p in self.doc.paragraphs:
            if regex.search(p.text):
                inline = p.runs
                # Loop added to work with runs (strings with same style)
                for i in range(len(inline)):
                    if regex.search(inline[i].text):
                        text = regex.sub(replace, inline[i].text)
                        inline[i].text = text
    
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    self.replace_regex(cell, regex , replace)
    
    def replace_multi(self,dictionary):
        for word, replacement in dictionary.items():
            word_re=re.compile(word)
            self.replace_regex(word_re, replacement)
        
import os

curdir = os.getcwd()
filename = 'simohua-twopage.docx'#'blank.docx'#
filepath_in = os.path.join(curdir,filename)
filepath_out = os.path.join(curdir,'out.docx')

def main():
    doc = DOC(filepath_in)
    doc.processtext()
    doc.save(filepath_out)
    
if __name__ == "__main__":
    main()